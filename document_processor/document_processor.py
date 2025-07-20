# document_processor.py
"""Core document processing logic with enhanced format detection."""

import logging
from pathlib import Path
from typing import List, Optional, Dict, Set, Tuple
import re
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml.ns import qn
from models import Document, DocumentSection
from utils import clean_text, truncate_text, extract_document_language

class DocumentProcessor:
    """Processes Word documents and extracts structured information."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.doc_counter = 0
        
        # Define proper heading styles
        self.proper_heading_styles = {
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 
            'Heading 5', 'Heading 6', 'Title', 'Subtitle'
        }
        
        # Define Arabic/English reference patterns
        self.arabic_reference_patterns = [
            r'المراجع',
            r'المصادر',
            r'قائمة المراجع',
            r'المصادر والمراجع'
        ]
        
        self.english_reference_patterns = [
            r'References',
            r'Bibliography',
            r'Works Cited',
            r'Sources'
        ]
        
        # Define Arabic image/table patterns for format checking
        self.arabic_figure_patterns = [
            r'\[الشكل\s*\d+\]',     # [الشكل 1]
            r'\[الصورة\s*\d+\]',    # [الصورة 1]
            r'\[الجدول\s*\d+\]',   # [الجدول 1]
            r'الشكل\s*\(\d+\)',     # الشكل (1)
            r'شكل\s*\d+',           # شكل 1
        ]
        
        # Caption patterns
        self.caption_patterns = [
            r'^\s*(?:الشكل|الصورة|الجدول|Figure|Table|Image)\s*\d+',
            r'^\s*\[.*\]\s*:?\s*.+',  # [Figure 1] Caption text
        ]
    
    def process_document(self, file_path: Path) -> Document:
        """Process a single Word document."""
        self.doc_counter += 1
        doc_name = file_path.name
        parent_folder = file_path.parent.name
        
        self.logger.info(f"Processing document {self.doc_counter}: {doc_name}")
        
        # Initialize document
        document = Document(
            id=self.doc_counter,
            file_path=file_path,
            name=doc_name,
            parent_folder=parent_folder,
            title="",
            author="Unknown",
            word_count=0,
            image_count=0,
            unique_image_count=0,
            author_from_text="Unknown",
            uses_proper_styles=True,
            arabic_reference_count=0,
            english_reference_count=0,
            footnote_count=0
        )
        
        # Add format tracking attributes
        document.format_issues = []
        document.heading_stats = {
            'total_headings': 0,
            'proper_style_headings': 0,
            'normal_style_headings': 0,
            'font_based_headings': 0
        }
        document.images_missing_captions = []
        document.heading_hierarchy_issues = []
        
        try:
            # Load document
            doc = DocxDocument(file_path)
            
            # Extract metadata (including author from properties)
            self._extract_metadata(doc, document)
            
            # Process paragraphs
            self._process_paragraphs(doc, document)
            
            # Count footnotes
            self._count_footnotes(doc, document)
            
            # Count unique images using the older, working method
            self._count_unique_images(doc, document)
            
            # Calculate format quality
            self._assess_format_quality(document)
            
        except Exception as e:
            self.logger.error(f"Error processing {doc_name}: {str(e)}")
            document.format_issues.append(f"Failed to process: {str(e)}")
        
        return document
    
    def _extract_metadata(self, doc: DocxDocument, document: Document):
        """Extract document metadata."""
        # Get core properties for author
        core_props = doc.core_properties
        
        # Extract author from properties
        if core_props.author:
            document.author = core_props.author
        elif core_props.creator:
            document.author = core_props.creator
        
        # Extract title
        if core_props.title:
            document.title = core_props.title
        else:
            # Get title from first paragraph
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                document.title = doc.paragraphs[0].text.strip()
        
        # Get author from second paragraph if exists
        if len(doc.paragraphs) >= 2 and doc.paragraphs[1].text.strip():
            second_para = doc.paragraphs[1].text.strip()
            if len(second_para) < 100:  # Reasonable length for author name
                document.author_from_text = second_para
    
    def _process_paragraphs(self, doc: DocxDocument, document: Document):
        """Process document paragraphs and sections."""
        full_text = []
        image_counter = 0
        seen_images = set()
        last_image_para_idx = -10  # Track last image position
        expected_caption_idx = -1
        in_references = False
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            full_text.append(text)
            
            # Handle first paragraph as Title (Heading 1)
            if para_idx == 0 and text:
                font_info = self._extract_font_info(paragraph)
                section = DocumentSection(
                    heading=text,
                    style_name=paragraph.style.name if paragraph.style else "Normal",
                    section_type="heading",
                    font_name=font_info.get('name'),
                    font_size=font_info.get('size'),
                    suggested_style="Heading 1"  # Title should be Heading 1
                )
                document.sections.append(section)
                
                # Update heading stats
                document.heading_stats['total_headings'] += 1
                
                # Track if it uses proper style
                if paragraph.style and paragraph.style.name in ['Title', 'Heading 1']:
                    document.heading_stats['proper_style_headings'] += 1
                else:
                    document.uses_proper_styles = False
                    document.heading_stats['normal_style_headings'] += 1
                    document.heading_hierarchy_issues.append({
                        'heading': text[:50],
                        'current_style': paragraph.style.name if paragraph.style else "Normal",
                        'suggested_style': "Heading 1"
                    })
                continue
            
            # Handle second paragraph as Author (keep as is, not a heading)
            if para_idx == 1 and text and len(text) < 100:
                document.author_from_text = text
                # Don't create a section for author - it's not a heading
                continue
            
            # Check if entering references section
            if self._is_reference_section(text):
                in_references = True
            
            # Count references
            if in_references:
                if self._is_arabic_reference(text):
                    document.arabic_reference_count += 1
                elif self._is_english_reference(text):
                    document.english_reference_count += 1
            
            # Check if it's an image/table reference
            if self._is_image_reference(text):
                image_counter += 1
                last_image_para_idx = para_idx
                expected_caption_idx = para_idx + 1
                
                # Create image section
                font_info = self._extract_font_info(paragraph)
                section = DocumentSection(
                    heading=text,
                    style_name=paragraph.style.name if paragraph.style else "Normal",
                    section_type="image",
                    font_name=font_info.get('name'),
                    font_size=font_info.get('size'),
                    suggested_style="Caption",
                    has_caption=False  # Will check next
                )
                document.sections.append(section)
                
            # Check if this could be a caption
            elif para_idx == expected_caption_idx and self._is_caption(text):
                # Mark previous image as having caption
                if document.sections and document.sections[-1].section_type == "image":
                    document.sections[-1].has_caption = True
                expected_caption_idx = -1
            
            # Check if it's a heading (max 20 words as per user requirement)
            elif self._is_heading(paragraph, text) and not in_references:
                # Determine heading level and suggested style
                current_style = paragraph.style.name if paragraph.style else "Normal"
                suggested_style = self._suggest_heading_style(text, current_style, document.sections)
                
                # Track heading style usage
                document.heading_stats['total_headings'] += 1
                if paragraph.style and paragraph.style.name in self.proper_heading_styles:
                    document.heading_stats['proper_style_headings'] += 1
                else:
                    document.uses_proper_styles = False
                    if paragraph.style and paragraph.style.name == "Normal":
                        document.heading_stats['normal_style_headings'] += 1
                        if self._has_special_formatting(paragraph):
                            document.heading_stats['font_based_headings'] += 1
                            document.format_issues.append(f"Heading using font size instead of style: {text[:50]}")
                
                # Create heading section
                font_info = self._extract_font_info(paragraph)
                section = DocumentSection(
                    heading=text,
                    style_name=current_style,
                    section_type="heading",
                    font_name=font_info.get('name'),
                    font_size=font_info.get('size'),
                    suggested_style=suggested_style
                )
                
                # Check if style matches suggestion
                if current_style != suggested_style and suggested_style != "Unknown":
                    document.heading_hierarchy_issues.append({
                        'heading': text[:50],
                        'current_style': current_style,
                        'suggested_style': suggested_style
                    })
                
                document.sections.append(section)
        
        # Check for images without captions
        for section in document.sections:
            if section.section_type == "image" and not section.has_caption:
                document.images_missing_captions.append(section.heading)
        
        # Set document statistics
        document.word_count = len(" ".join(full_text).split())
        document.image_count = image_counter
    
    def _is_heading(self, paragraph, text: str) -> bool:
        """Check if paragraph is a heading (max 20 words)."""
        # Skip if it looks like a reference
        if self._looks_like_reference(text):
            return False
        
        # Check word count - headings should be max 20 words
        word_count = len(text.split())
        if word_count > 20:
            return False
        
        # Check if using proper heading style
        if paragraph.style and paragraph.style.name in self.proper_heading_styles:
            return True
        
        # Check if it looks like a heading but uses wrong style
        if self._has_special_formatting(paragraph):
            # Additional checks to avoid false positives
            # Skip if it's just a number or very short
            if len(text) < 3 or text.isdigit():
                return False
            return True
        
        # Check for numbered headings (e.g., "1. Introduction", "2.1 Background")
        if re.match(r'^\d+(\.\d+)*\.?\s+\w+', text):
            return True
        
        # Check for bullet points that might be headings
        if text.startswith('•') and word_count <= 20:
            return True
        
        return False
    
    def _suggest_heading_style(self, text: str, current_style: str, existing_sections: List[DocumentSection]) -> str:
        """Suggest appropriate heading style based on document structure."""
        # If already using proper heading style, keep it
        if current_style in self.proper_heading_styles:
            return current_style
        
        # Count existing heading levels
        heading_counts = {'Heading 1': 0, 'Heading 2': 0, 'Heading 3': 0, 'Heading 4': 0}
        for section in existing_sections:
            if section.section_type == "heading" and section.style_name in heading_counts:
                heading_counts[section.style_name] += 1
        
        # Analyze text patterns
        # Numbered patterns like "1.", "1.1", "1.1.1"
        number_match = re.match(r'^(\d+)(\.(\d+))*(\.(\d+))*\.?\s', text)
        if number_match:
            dots = text.split('.')[:-1]  # Count dots
            level = len(dots)
            if level == 1:
                return "Heading 2"  # Main sections
            elif level == 2:
                return "Heading 3"  # Subsections
            elif level >= 3:
                return "Heading 4"  # Sub-subsections
        
        # Arabic numbering
        if re.match(r'^[أ-ي]\.', text) or re.match(r'^[١-٩]\.', text):
            return "Heading 2"
        
        # Bullet points
        if text.startswith('•'):
            return "Heading 3"
        
        # If it's short and bold/large, likely a main heading
        if len(text) < 50:
            return "Heading 2"
        
        # Default suggestion based on position
        if heading_counts['Heading 2'] == 0:
            return "Heading 2"
        elif heading_counts['Heading 3'] == 0:
            return "Heading 3"
        else:
            return "Heading 4"
    
    def _is_caption(self, text: str) -> bool:
        """Check if text is likely a caption."""
        for pattern in self.caption_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Also check if it starts with descriptive text after an image
        if len(text) < 200 and not text.endswith('.'):
            return True
        
        return False
    
    def _is_reference_section(self, text: str) -> bool:
        """Check if entering references section."""
        combined_patterns = self.arabic_reference_patterns + self.english_reference_patterns
        for pattern in combined_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def _is_arabic_reference(self, text: str) -> bool:
        """Check if text is an Arabic reference entry."""
        # Simple heuristic: contains Arabic text and has reference-like structure
        arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')
        has_arabic = bool(arabic_pattern.search(text))
        
        # Check for common reference patterns (year in parentheses, dots, etc.)
        has_year = bool(re.search(r'\(\d{4}\)', text))
        has_dots = text.count('.') >= 2
        
        return has_arabic and (has_year or has_dots)
    
    def _is_english_reference(self, text: str) -> bool:
        """Check if text is an English reference entry."""
        # Check for common reference patterns
        has_year = bool(re.search(r'\(\d{4}\)', text))
        has_dots = text.count('.') >= 2
        has_comma = ',' in text
        
        # Exclude if mostly Arabic
        arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')
        arabic_chars = len(arabic_pattern.findall(text))
        total_chars = len(text)
        
        is_mostly_english = arabic_chars < total_chars * 0.3
        
        return is_mostly_english and (has_year or (has_dots and has_comma))
    
    def _looks_like_reference(self, text: str) -> bool:
        """Check if text looks like a reference entry."""
        return self._is_arabic_reference(text) or self._is_english_reference(text)
    
    def _count_footnotes(self, doc: DocxDocument, document: Document):
        """Count footnotes in the document."""
        try:
            footnote_count = 0
            
            # Method 1: Check document part relationships for footnotes
            for rel in doc.part.rels.values():
                if "footnotes" in rel.reltype:
                    # Found footnotes relationship
                    self.logger.info("Found footnotes relationship")
                    
                    # Try to access footnotes part
                    try:
                        footnotes_part = rel.target_part
                        # Count footnote elements (usually starts from id=1, id=0 is separator)
                        footnote_elements = footnotes_part.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote')
                        # Subtract separator footnotes (usually first 2)
                        footnote_count = max(0, len(footnote_elements) - 2)
                        self.logger.info(f"Counted {len(footnote_elements)} footnote elements, {footnote_count} actual footnotes")
                    except Exception as e:
                        self.logger.warning(f"Could not access footnotes part: {e}")
            
            # Method 2: Count footnote references in the document body
            if footnote_count == 0:
                # Look for footnote references in paragraphs
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        # Check for footnote references
                        footnote_refs = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnoteReference')
                        footnote_count += len(footnote_refs)
                
                if footnote_count > 0:
                    self.logger.info(f"Found {footnote_count} footnote references in document body")
            
            document.footnote_count = footnote_count
            
            if footnote_count > 0:
                self.logger.info(f"Total footnotes found: {footnote_count}")
            
        except Exception as e:
            self.logger.warning(f"Could not count footnotes: {e}")
            document.footnote_count = 0
    
    def _count_unique_images(self, doc: DocxDocument, document: Document):
        """Count unique images by detecting PNG/SVG duplicate patterns."""
        # Collect all image references
        all_images = {}
        
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                filename = rel.target_ref.split('/')[-1]
                all_images[filename] = rel_id
        
        if not all_images:
            document.unique_image_count = 0
            return
        
        self.logger.info(f"====================Processing: {document.file_path.name}====================")
        self.logger.info(f"  Found images: {sorted(all_images.keys())}")
        
        # Identify SVG duplicates to ignore
        images_to_ignore = set()
        
        for filename in all_images:
            if filename.endswith('.svg'):
                match = re.search(r'image(\d+)\.svg', filename)
                if match:
                    svg_num = int(match.group(1))
                    png_equivalent = f"image{svg_num - 1}.png"
                    
                    if png_equivalent in all_images:
                        images_to_ignore.add(filename)
                        self.logger.info(f"  Ignoring {filename} (duplicate of {png_equivalent})")
        
        unique_images = set(all_images.keys()) - images_to_ignore
        unique_count = len(unique_images)
        
        self.logger.info(f"  Unique images: {sorted(unique_images)}")
        self.logger.info(f"  Total files: {len(all_images)}, Unique images: {unique_count}")
        
        document.unique_image_count = unique_count
    
    def _has_special_formatting(self, paragraph) -> bool:
        """Check if paragraph has special formatting (bold, large font, etc.)."""
        if not paragraph.runs:
            return False
        
        # Check each run in the paragraph
        for run in paragraph.runs:
            # Check for bold
            if run.font.bold:
                return True
            
            # Check for large font size (14pt or larger)
            if run.font.size and run.font.size.pt >= 14:
                return True
        
        return False
    
    def _extract_font_info(self, paragraph) -> Dict[str, any]:
        """Extract font information from paragraph."""
        font_info = {'name': None, 'size': None}
        
        if paragraph.runs:
            # Get font from first run
            run = paragraph.runs[0]
            if run.font.name:
                font_info['name'] = run.font.name
            if run.font.size:
                font_info['size'] = run.font.size.pt
        
        return font_info
    
    def _is_image_reference(self, text: str) -> bool:
        """Check if text is an image/table reference."""
        # Check Arabic patterns
        for pattern in self.arabic_figure_patterns:
            if re.search(pattern, text):
                return True
        
        # Check English patterns
        english_patterns = [
            r'\[Figure\s*\d+\]',
            r'\[Table\s*\d+\]',
            r'Figure\s*\d+',
            r'Table\s*\d+',
        ]
        
        for pattern in english_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _assess_format_quality(self, document: Document):
        """Assess overall document format quality."""
        # Calculate quality metrics
        if document.heading_stats['total_headings'] > 0:
            proper_ratio = document.heading_stats['proper_style_headings'] / document.heading_stats['total_headings']
        else:
            proper_ratio = 1.0  # No headings is not necessarily bad
        
        # Count issues
        issue_count = len(document.format_issues) + len(document.images_missing_captions) + len(document.heading_hierarchy_issues)
        
        # Determine quality
        if proper_ratio >= 0.9 and issue_count == 0:
            document.format_quality = 'Excellent'
        elif proper_ratio >= 0.7 and issue_count <= 3:
            document.format_quality = 'Good'
        elif proper_ratio >= 0.5 or issue_count <= 10:
            document.format_quality = 'Fair'
        else:
            document.format_quality = 'Poor'
        
        # Log format summary
        if document.format_quality in ['Poor', 'Fair']:
            self.logger.warning(f"Document '{document.name}' has formatting issues:")
            self.logger.warning(f"  - Proper heading usage: {proper_ratio:.0%}")
            self.logger.warning(f"  - Total issues: {issue_count}")
            if document.heading_stats['normal_style_headings'] > 0:
                self.logger.warning(f"  - {document.heading_stats['normal_style_headings']} headings using Normal style")
            if document.images_missing_captions:
                self.logger.warning(f"  - {len(document.images_missing_captions)} images missing captions")
    
    def process_folder(self, folder_path: Path) -> List[Document]:
        """Process all documents in a folder."""
        documents = []
        
        # Find all .docx files
        docx_files = list(folder_path.rglob("*.docx"))
        
        # Filter out temporary files
        docx_files = [f for f in docx_files if not f.name.startswith("~")]
        
        self.logger.info(f"Found {len(docx_files)} documents to process")
        
        # Process each document
        for file_path in docx_files:
            try:
                doc = self.process_document(file_path)
                documents.append(doc)
            except Exception as e:
                self.logger.error(f"Failed to process {file_path}: {e}")
        
        # Log summary statistics
        self._log_processing_summary(documents)
        
        return documents
    
    def _log_processing_summary(self, documents: List[Document]):
        """Log summary of processing results."""
        self.logger.info("\n" + "=" * 60)
        self.logger.info("DOCUMENT PROCESSING SUMMARY")
        self.logger.info("=" * 60)
        
        # Count documents by quality
        quality_counts = {}
        for doc in documents:
            quality = getattr(doc, 'format_quality', 'Unknown')
            quality_counts[quality] = quality_counts.get(quality, 0) + 1
        
        # Log quality distribution
        self.logger.info("Format Quality Distribution:")
        for quality in ['Excellent', 'Good', 'Fair', 'Poor']:
            count = quality_counts.get(quality, 0)
            if count > 0:
                self.logger.info(f"  {quality}: {count} documents ({count/len(documents)*100:.1f}%)")
        
        # Log documents with issues
        poor_docs = [doc for doc in documents if getattr(doc, 'format_quality', '') in ['Poor', 'Fair']]
        if poor_docs:
            self.logger.info(f"\nDocuments needing attention ({len(poor_docs)}):")
            for doc in poor_docs[:10]:  # Show first 10
                self.logger.info(f"  - {doc.name}")