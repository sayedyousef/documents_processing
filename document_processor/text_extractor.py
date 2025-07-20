# text_extractor.py
"""Module for extracting text-based information from documents."""

from docx import Document as DocxDocument

class TextExtractor:
    """Handles extraction of text content from documents."""
    
    def extract_title(self, docx: DocxDocument) -> str:
        """Extract document title from first non-empty line."""
        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if text:  # First non-empty line is the title
                return text
        return "Untitled"
    
    def extract_author(self, docx: DocxDocument) -> str:
        """Extract document author from second non-empty line."""
        non_empty_count = 0
        
        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if text:
                non_empty_count += 1
                if non_empty_count == 2:  # Second non-empty line is the author
                    return text
        return "Unknown Author"
    
    def count_words(self, docx: DocxDocument) -> int:
        """Count total words in document."""
        word_count = 0
        
        # Count words in paragraphs
        for paragraph in docx.paragraphs:
            if paragraph.text:
                word_count += len(paragraph.text.split())
        
        # Count words in tables
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            word_count += len(paragraph.text.split())
        
        return word_count