# document_processor.py
"""Core document processing functionality."""

import logging
from pathlib import Path
from typing import List
from docx import Document as DocxDocument
from docx.shared import Pt
from models import Document, DocumentSection
from utils import clean_text, is_arabic_text

class DocumentProcessor:
    """Handles processing of Word documents."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.doc_counter = 0
        
    def process_folder(self, folder_path: Path) -> List[Document]:
        """Process all Word documents in folder and subfolders."""
        documents = []
        
        # Find all .docx files recursively
        for doc_path in folder_path.rglob("*.docx"):
            # Skip temporary files
            if doc_path.name.startswith("~"):
                continue
                
            try:
                self.logger.info(f"Processing: {doc_path}")
                doc = self.process_document(doc_path)
                documents.append(doc)
            except Exception as e:
                self.logger.error(f"Error processing {doc_path}: {e}")
                
        return documents
    
    def process_document(self, file_path: Path) -> Document:
        """Process a single Word document."""
        self.doc_counter += 1
        docx = DocxDocument(file_path)
        
        # Extract basic information
        doc = Document(
            id=self.doc_counter,
            file_path=file_path,
            name=file_path.name,
            parent_folder=file_path.parent.name,
            title=self._extract_title(docx),
            word_count=self._count_words(docx),
            image_count=self._count_images_total(docx),
            unique_image_count=self._count_unique_images(docx, file_path),
            sections=self._extract_sections(docx)
        )
        
        return doc
    
    def _extract_title(self, docx: DocxDocument) -> str:
        """Extract document title from core properties or first heading."""
        # Try to get from document properties
        if docx.core_properties.title:
            return docx.core_properties.title
        
        # Otherwise, look for first heading
        for paragraph in docx.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                return paragraph.text.strip()
        
        return "Untitled"
    
    def _count_words(self, docx: DocxDocument) -> int:
        """Count total words in document."""
        word_count = 0
        
        # Count words in paragraphs
        for paragraph in docx.paragraphs:
            if paragraph.text:
                word_count += len(paragraph.text.split())
        
        # Count words in tables
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            word_count += len(paragraph.text.split())
        
        return word_count

    def _count_images_total(self, docx: DocxDocument) -> int:
        """Count total images in document (including all formats/duplicates)."""
        # Use a set to track unique image relationships
        unique_images = set()
        
        # Method 1: Check document relationships
        for rel_id, rel in docx.part.rels.items():
            if "image" in rel.reltype:
                unique_images.add(rel.target_ref)
        
        # Method 2: Check inline shapes (more reliable)
        try:
            # Register namespace for picture elements
            from lxml import etree
            nsmap = {'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}
            
            # Count inline pictures in paragraphs
            for paragraph in docx.paragraphs:
                for pic in paragraph._element.xpath('.//pic:pic', namespaces=nsmap):
                    # Try to get relationship ID
                    blip = pic.xpath('.//a:blip/@r:embed', 
                                   namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                                             'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
                    if blip:
                        unique_images.add(blip[0])
        except Exception as e:
            self.logger.warning(f"Error counting inline images: {e}")
        
        # Count all unique image references
        return len(unique_images)

    def _count_unique_images(self, docx: DocxDocument, file_path: Path) -> int:
        """Count unique images in document (grouped by base name)."""
        unique_images = set()
        
        # Get actual image filenames from relationships
        for rel_id, rel in docx.part.rels.items():
            if "image" in rel.reltype:
                # Get the actual target path (e.g., "media/image2.png")
                target_path = rel.target_ref
                unique_images.add(target_path)
        
        # If no images found via relationships, return 0
        if not unique_images:
            return 0
        
        # Log document being processed
        self.logger.info(f"====================Processing: {file_path.name}====================")

        # Group images by base name (without extension)
        image_groups = {}
        for img_path in unique_images:
            # Extract just the filename from path like "media/image2.png" → "image2"
            filename = img_path.split('/')[-1]  # Get filename from path
            base_name = filename.split('.')[0]  # Remove extension
            
            # Log each image found
            self.logger.info(f"  Found image: {img_path} (base: {base_name})")
      
            if base_name not in image_groups:
                image_groups[base_name] = []
            image_groups[base_name].append(img_path)
        
        # Log summary
        self.logger.info(f"  Image groups: {list(image_groups.keys())}")
        self.logger.info(f"  Total files: {len(unique_images)}, Unique images: {len(image_groups)}")
        
        # Count unique images (one per base name)
        return len(image_groups)
    
    def _extract_sections(self, docx: DocxDocument) -> List[DocumentSection]:
        """Extract sections with their fonts and content."""
        sections = []
        current_section = None
        
        for paragraph in docx.paragraphs:
            # Check if this is a heading
            if paragraph.style.name.startswith('Heading') or self._has_special_formatting(paragraph):
                # Save previous section if exists
                if current_section:
                    sections.append(current_section)
                
                # Get font information
                font_info = self._get_font_info(paragraph)
                
                # Create new section
                current_section = DocumentSection(
                    heading=clean_text(paragraph.text),
                    font_name=font_info['name'],
                    font_size=font_info['size']
                )
            elif current_section and paragraph.text.strip():
                # Add text to current section
                current_section.text += paragraph.text + "\n"
        
        # Don't forget the last section
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _has_special_formatting(self, paragraph) -> bool:
        """Check if paragraph has special formatting (bold, larger font)."""
        if not paragraph.runs:
            return False
        
        # Check first run for bold or larger font
        first_run = paragraph.runs[0]
        if first_run.bold:
            return True
        
        if first_run.font.size and first_run.font.size > Pt(14):
            return True
        
        return False
    
    def _get_font_info(self, paragraph) -> dict:
        """Extract font information from paragraph."""
        font_info = {'name': None, 'size': None}
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            
            # Get font name
            if first_run.font.name:
                font_info['name'] = first_run.font.name
            elif paragraph.style.font.name:
                font_info['name'] = paragraph.style.font.name
            
            # Get font size
            if first_run.font.size:
                font_info['size'] = first_run.font.size.pt
            elif paragraph.style.font.size:
                font_info['size'] = paragraph.style.font.size.pt
        
        return font_info